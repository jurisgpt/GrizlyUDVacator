from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt


def create_template():
    # Create a new document
    doc = Document()

    # Add title
    title = doc.add_heading("MOTION TO VACATE DEFAULT JUDGMENT", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add subtitle with court information
    subtitle = doc.add_paragraph()
    subtitle.add_run("Superior Court of California, County of Los Angeles").bold = True
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add case information section
    doc.add_heading("Case Information", 1)
    case_info = doc.add_paragraph()
    case_info.add_run("Case Number: {{ case_number }}\n")
    case_info.add_run("Plaintiff: {{ plaintiff }}\n")
    case_info.add_run("Defendant: {{ defendant }}")

    # Add motion section
    doc.add_heading("Motion", 1)
    motion = doc.add_paragraph()
    motion.add_run("Dear {{ court_name }},\n\n")
    motion.add_run(
        "I, {{ defendant }}, respectfully request relief under CCP § 473(b) and CCP § 473.5.\n\n"
    )

    # Add justification section
    doc.add_heading("Justification", 1)
    justification = doc.add_paragraph()
    justification.add_run("Pursuant to {{ statute_list }}:\n")

    # Add facts section
    doc.add_heading("Facts", 1)
    facts = doc.add_paragraph()
    facts.add_run("The following facts support this motion:\n")

    # Add signature block
    doc.add_heading("Signature", 1)
    signature = doc.add_paragraph()
    signature.add_run("Respectfully submitted,\n\n\n")
    signature.add_run("{{ defendant }}")

    # Add footer
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.add_run("Date: {{ current_date }}")
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Save the document
    doc.save("backend/generator/templates/motion_template.docx")


if __name__ == "__main__":
    create_template()
